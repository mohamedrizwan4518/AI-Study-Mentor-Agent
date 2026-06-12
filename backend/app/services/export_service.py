import io
import re
import logging
import tempfile
import os
from pathlib import Path
from markdown2 import markdown
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle, Preformatted, Image
from reportlab.lib import colors
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import matplotlib
matplotlib.use('Agg')  # Use non-interactive backend
import matplotlib.pyplot as plt
from matplotlib import mathtext

# ==========================
# Configuration
# ==========================
WATERMARK_TEXT = "~honeypot"

logger = logging.getLogger(__name__)

# ==========================
# Unicode Font Registration
# ==========================

def _register_unicode_fonts() -> str:
    """
    Register a Unicode-capable TrueType font family with ReportLab.

    Tries DejaVuSans (common on Linux/Docker images) first, then falls back
    to the Vera font bundled with ReportLab, and finally to Helvetica.

    Returns the base font name that was registered (e.g. 'DejaVuSans').
    """
    candidates = [
        {
            "name": "DejaVuSans",
            "regular": [
                "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
                "/usr/share/fonts/dejavu/DejaVuSans.ttf",
                "/usr/share/fonts/truetype/ttf-dejavu/DejaVuSans.ttf",
            ],
            "bold": [
                "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
                "/usr/share/fonts/dejavu/DejaVuSans-Bold.ttf",
                "/usr/share/fonts/truetype/ttf-dejavu/DejaVuSans-Bold.ttf",
            ],
            "italic": [
                "/usr/share/fonts/truetype/dejavu/DejaVuSans-Oblique.ttf",
                "/usr/share/fonts/dejavu/DejaVuSans-Oblique.ttf",
            ],
            "boldItalic": [
                "/usr/share/fonts/truetype/dejavu/DejaVuSans-BoldOblique.ttf",
                "/usr/share/fonts/dejavu/DejaVuSans-BoldOblique.ttf",
            ],
        },
    ]

    import reportlab as _rl
    _rl_fonts_dir = os.path.join(os.path.dirname(_rl.__file__), "fonts")

    for candidate in candidates:
        name = candidate["name"]
        paths = {}
        ok = True
        for variant, path_list in candidate.items():
            if variant == "name":
                continue
            found = next((p for p in path_list if os.path.exists(p)), None)
            if found is None:
                ok = False
                break
            paths[variant] = found

        if not ok:
            continue

        try:
            pdfmetrics.registerFont(TTFont(name, paths["regular"]))
            pdfmetrics.registerFont(TTFont(f"{name}-Bold", paths["bold"]))
            pdfmetrics.registerFont(TTFont(f"{name}-Italic", paths["italic"]))
            pdfmetrics.registerFont(TTFont(f"{name}-BoldItalic", paths["boldItalic"]))
            pdfmetrics.registerFontFamily(
                name,
                normal=name,
                bold=f"{name}-Bold",
                italic=f"{name}-Italic",
                boldItalic=f"{name}-BoldItalic",
            )
            logger.debug(f"Registered Unicode font family: {name}")
            return name
        except Exception as exc:
            logger.warning(f"Could not register {name}: {exc}")

    # Fallback: Vera (bundled with ReportLab) – supports more Unicode than Helvetica
    vera_regular = os.path.join(_rl_fonts_dir, "Vera.ttf")
    vera_bold = os.path.join(_rl_fonts_dir, "VeraBd.ttf")
    vera_italic = os.path.join(_rl_fonts_dir, "VeraIt.ttf")
    vera_bolditalic = os.path.join(_rl_fonts_dir, "VeraBI.ttf")
    if all(os.path.exists(p) for p in [vera_regular, vera_bold, vera_italic, vera_bolditalic]):
        try:
            pdfmetrics.registerFont(TTFont("Vera", vera_regular))
            pdfmetrics.registerFont(TTFont("Vera-Bold", vera_bold))
            pdfmetrics.registerFont(TTFont("Vera-Italic", vera_italic))
            pdfmetrics.registerFont(TTFont("Vera-BoldItalic", vera_bolditalic))
            pdfmetrics.registerFontFamily(
                "Vera",
                normal="Vera",
                bold="Vera-Bold",
                italic="Vera-Italic",
                boldItalic="Vera-BoldItalic",
            )
            logger.debug("Registered Unicode font family: Vera (fallback)")
            return "Vera"
        except Exception as exc:
            logger.warning(f"Could not register Vera: {exc}")

    logger.warning("No Unicode TTF font available; falling back to Helvetica (limited Unicode support)")
    return "Helvetica"


# Register fonts at module load time so styles can reference the family name.
_UNICODE_FONT = _register_unicode_fonts()

# ==========================
# ReportLab PDF Utilities
# ==========================

class WatermarkCanvas(canvas.Canvas):
    """Custom canvas class to add watermark to PDF pages."""
    def __init__(self, *args, **kwargs):
        self.add_watermark = kwargs.pop('add_watermark', False)
        canvas.Canvas.__init__(self, *args, **kwargs)
    
    def showPage(self):
        """Override to add watermark before showing page."""
        if self.add_watermark:
            self.saveState()
            self.setFont('Helvetica', 10)
            self.setFillColor(colors.lightgrey)
            # Add watermark at bottom right
            self.drawRightString(
                A4[0] - 20,
                20,
                WATERMARK_TEXT
            )
            self.restoreState()
        canvas.Canvas.showPage(self)

def render_latex_to_image(latex_text: str, inline: bool = True) -> str:
    """
    Render LaTeX formula to a PNG image and return the file path.
    
    Args:
        latex_text: The LaTeX formula (without $ delimiters)
        inline: Whether this is inline math (True) or display math (False)
    
    Returns:
        Path to the generated PNG image file
    """
    try:
        # Create a temporary file for the image
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
        temp_file.close()
        
        # Use larger font size for better readability
        fontsize = 20 if inline else 24
        text = f'${latex_text}$'
        
        # Create a figure with transparent background for size calculation
        fig, ax = plt.subplots(figsize=(1, 1))
        ax.axis('off')
        fig.patch.set_visible(False)
        
        # Render text to calculate bounding box
        t = ax.text(0.5, 0.5, text, fontsize=fontsize, ha='center', va='center')
        
        # Draw to calculate proper bbox
        fig.canvas.draw()
        bbox = t.get_window_extent(renderer=fig.canvas.get_renderer())
        
        # Convert to figure coordinates
        bbox_inches = bbox.transformed(fig.dpi_scale_trans.inverted())
        
        plt.close(fig)
        
        # Add padding (in inches)
        padding = 0.2
        width = bbox_inches.width + padding
        height = bbox_inches.height + padding
        
        # Ensure minimum size
        width = max(width, 1.0)
        height = max(height, 0.4)
        
        # Create final figure with calculated size
        dpi = 200  # High DPI for quality
        fig = plt.figure(figsize=(width, height), dpi=dpi)
        fig.patch.set_facecolor('white')
        
        # Add axes that fills the figure
        ax = fig.add_axes([0, 0, 1, 1])
        ax.axis('off')
        ax.set_xlim(0, 1)
        ax.set_ylim(0, 1)
        
        # Render the LaTeX text centered
        ax.text(0.5, 0.5, text,
                fontsize=fontsize,
                ha='center',
                va='center',
                color='black')
        
        # Save with high quality
        plt.savefig(temp_file.name,
                   format='png',
                   dpi=dpi,
                   bbox_inches='tight',
                   pad_inches=0.15,
                   facecolor='white',
                   edgecolor='none')
        plt.close(fig)
        
        logger.debug(f"Rendered LaTeX to image: {temp_file.name} (size: {width}x{height} inches @ {dpi} DPI)")
        return temp_file.name
        
    except Exception as e:
        logger.error(f"Failed to render LaTeX '{latex_text}': {str(e)}")
        return None

def create_styles():
    """Create custom paragraph styles for PDF."""
    styles = getSampleStyleSheet()

    # Derive bold/italic font names from the registered Unicode font family.
    # When _UNICODE_FONT is 'Helvetica' the built-in Helvetica-Bold/-Oblique
    # names are used; for TrueType families the registered variant names are used.
    if _UNICODE_FONT == "Helvetica":
        _bold = "Helvetica-Bold"
        _italic = "Helvetica-Oblique"
    else:
        _bold = f"{_UNICODE_FONT}-Bold"
        _italic = f"{_UNICODE_FONT}-Italic"

    # Title style
    styles.add(ParagraphStyle(
        name='CustomTitle',
        parent=styles['Title'],
        fontSize=28,
        textColor=colors.HexColor('#1a1a1a'),
        spaceAfter=30,
        alignment=TA_CENTER,
        fontName=_bold
    ))
    
    # Heading 2 style
    styles.add(ParagraphStyle(
        name='CustomHeading2',
        parent=styles['Heading2'],
        fontSize=20,
        textColor=colors.HexColor('#2a2a2a'),
        spaceAfter=10,
        spaceBefore=20,
        fontName=_bold,
        borderWidth=0,
        borderPadding=5,
        borderColor=colors.HexColor('#e0e0e0'),
        borderRadius=0,
    ))
    
    # Heading 3 style
    styles.add(ParagraphStyle(
        name='CustomHeading3',
        parent=styles['Heading3'],
        fontSize=16,
        textColor=colors.HexColor('#3a3a3a'),
        spaceAfter=8,
        spaceBefore=15,
        fontName=_bold
    ))
    
    # Body text style
    styles.add(ParagraphStyle(
        name='CustomBody',
        parent=styles['BodyText'],
        fontSize=11,
        textColor=colors.HexColor('#333333'),
        alignment=TA_JUSTIFY,
        spaceAfter=10,
        fontName=_UNICODE_FONT
    ))
    
    # Code style
    styles.add(ParagraphStyle(
        name='CustomCode',
        parent=styles['Code'],
        fontSize=9,
        fontName='Courier',
        textColor=colors.black,
        backColor=colors.HexColor('#f4f4f4'),
        leftIndent=10,
        rightIndent=10,
        spaceAfter=10,
        spaceBefore=10
    ))
    
    return styles


def normalize_heading_level(line: str) -> tuple:
    """
    Normalize heading levels: treat 4+ '#' symbols as level 3 (###).
    
    Args:
        line: The markdown line to check
        
    Returns:
        Tuple of (normalized_line, level) where level is 0 (not a heading), 1, 2, or 3
    """
    match = re.match(r'^(#{1,})\s+(.+)$', line)
    if not match:
        return line, 0
    
    hashes, text = match.groups()
    level = len(hashes)
    
    # Normalize: treat 4+ as level 3
    if level >= 4:
        level = 3
        # Return normalized line with exactly 3 hashes
        return f"### {text}", level
    
    return line, level


def parse_markdown_to_reportlab(content: str, styles) -> list:
    """
    Parse markdown content and convert to ReportLab flowables.
    Supports headings, paragraphs, bold, italic, code blocks, lists, and LaTeX formulas.
    LaTeX formulas are rendered as images and placed inline.

    Heading normalization: Any heading with 4 or more '#' symbols is treated as '###'.
    """
    # Pre-process the full content to normalise multi-line LaTeX delimiters
    # ($$...$$, \\[...\\], \\(...\\)) before splitting into lines.  This is
    # necessary because those environments can span multiple lines and would
    # otherwise be missed by the per-line format_inline_markdown processing.
    content = fix_latex_delimiters(content)

    # Normalise line endings: strip \r so that \r\n becomes \n.
    # This MUST happen before splitting into lines, otherwise trailing \r
    # characters survive and render as ■ in PDF fonts (especially Courier
    # used in code blocks, which bypasses sanitize_for_xml).
    content = content.replace('\r', '')

    # Strip problematic Unicode characters that PDF fonts can't render (shown
    # as ■ black squares).  These commonly appear in AI-generated markdown.
    _INVISIBLE_CHARS = str.maketrans({
        '\u200b': '',    # zero-width space
        '\u200c': '',    # zero-width non-joiner
        '\u200d': '',    # zero-width joiner
        '\u200e': '',    # left-to-right mark
        '\u200f': '',    # right-to-left mark
        '\u2028': '\n',  # line separator → newline
        '\u2029': '\n',  # paragraph separator → newline
        '\ufeff': '',    # BOM / zero-width no-break space
        '\u00ad': '',    # soft hyphen
        '\u2060': '',    # word joiner
        '\ufffe': '',    # non-character
        '\uffff': '',    # non-character
    })
    content = content.translate(_INVISIBLE_CHARS)

    elements = []
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    in_list = False
    list_items = []
    temp_image_files = []  # Track temporary image files for cleanup
    
    while i < len(lines):
        line = lines[i]
        
        # Handle code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End of code block
                code_text = '\n'.join(code_lines)
                elements.append(Preformatted(code_text, styles['CustomCode']))
                elements.append(Spacer(1, 0.2*inch))
                code_lines = []
                in_code_block = False
            else:
                # Start of code block
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        # Normalize heading levels (treat 4+ '#' as 3)
        normalized_line, heading_level = normalize_heading_level(line)
        
        # Handle headings
        if heading_level == 1:
            if in_list:
                elements.extend(list_items)
                list_items = []
                in_list = False
            text = normalized_line[2:].strip()  # Remove '# '
            processed_elements = format_inline_markdown(text, styles, temp_image_files)
            if isinstance(processed_elements, list):
                elements.extend(processed_elements)
            else:
                elements.append(Paragraph(processed_elements, styles['CustomTitle']))
            elements.append(Spacer(1, 0.3*inch))
        elif heading_level == 2:
            if in_list:
                elements.extend(list_items)
                list_items = []
                in_list = False
            text = normalized_line[3:].strip()  # Remove '## '
            processed_elements = format_inline_markdown(text, styles, temp_image_files)
            if isinstance(processed_elements, list):
                elements.extend(processed_elements)
            else:
                elements.append(Paragraph(processed_elements, styles['CustomHeading2']))
            elements.append(Spacer(1, 0.2*inch))
        elif heading_level == 3:
            if in_list:
                elements.extend(list_items)
                list_items = []
                in_list = False
            text = normalized_line[4:].strip()  # Remove '### '
            processed_elements = format_inline_markdown(text, styles, temp_image_files)
            if isinstance(processed_elements, list):
                elements.extend(processed_elements)
            else:
                elements.append(Paragraph(processed_elements, styles['CustomHeading3']))
            elements.append(Spacer(1, 0.15*inch))
        # Handle lists
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:].strip()
            processed_elements = format_inline_markdown(text, styles, temp_image_files)
            if isinstance(processed_elements, list):
                # If there are images, add them with proper bullet handling
                first_text_added = False
                for elem in processed_elements:
                    if isinstance(elem, Image):
                        list_items.append(elem)
                    elif isinstance(elem, Paragraph):
                        # Add bullet to the first text paragraph only
                        if not first_text_added:
                            # Extract text and prepend bullet
                            elem_text = elem.text if hasattr(elem, 'text') else str(elem)
                            list_items.append(Paragraph(f"• {elem_text}", elem.style))
                            first_text_added = True
                        else:
                            list_items.append(elem)
            else:
                list_items.append(Paragraph(f"• {processed_elements}", styles['CustomBody']))
            in_list = True
        elif re.match(r'^\d+\.\s', line.strip()):
            # Numbered list (supports any number)
            match = re.match(r'^(\d+)\.\s+(.+)', line.strip())
            if match:
                num, text = match.groups()
                processed_elements = format_inline_markdown(text, styles, temp_image_files)
                if isinstance(processed_elements, list):
                    # If there are images, add them with proper numbering handling
                    first_text_added = False
                    for elem in processed_elements:
                        if isinstance(elem, Image):
                            list_items.append(elem)
                        elif isinstance(elem, Paragraph):
                            # Add number to the first text paragraph only
                            if not first_text_added:
                                # Extract text and prepend number
                                elem_text = elem.text if hasattr(elem, 'text') else str(elem)
                                list_items.append(Paragraph(f"{num}. {elem_text}", elem.style))
                                first_text_added = True
                            else:
                                list_items.append(elem)
                else:
                    list_items.append(Paragraph(f"{num}. {processed_elements}", styles['CustomBody']))
                in_list = True
        # Handle horizontal rules
        elif line.strip() in ['---', '***', '___']:
            if in_list:
                elements.extend(list_items)
                list_items = []
                in_list = False
            elements.append(Spacer(1, 0.2*inch))
        # Handle empty lines
        elif not line.strip():
            if in_list:
                elements.extend(list_items)
                elements.append(Spacer(1, 0.1*inch))
                list_items = []
                in_list = False
            elif elements:  # Don't add spacer at the beginning
                elements.append(Spacer(1, 0.1*inch))
        # Handle regular paragraphs
        else:
            if in_list:
                elements.extend(list_items)
                list_items = []
                in_list = False
            processed_elements = format_inline_markdown(line, styles, temp_image_files)
            if isinstance(processed_elements, list):
                # Mix of text and images - add them all
                elements.extend(processed_elements)
            elif processed_elements and processed_elements.strip():
                elements.append(Paragraph(processed_elements, styles['CustomBody']))
        
        i += 1
    
    # Add any remaining list items
    if in_list:
        elements.extend(list_items)
    
    return elements, temp_image_files


def fix_latex_delimiters(text: str) -> str:
    """
    Fix common LaTeX delimiter issues and normalise all LaTeX delimiters to
    the single-dollar ``$...$`` form expected by the rendering pipeline.

    Conversions performed (in order):
    1. Backtick-enclosed LaTeX: `` `$...$` `` or `` `$...` `` → ``$...$``
    2. Display-math double-dollars: ``$$...$$`` → ``$...$``
    3. LaTeX display-math brackets: ``\\[...\\]`` → ``$...$``
    4. LaTeX inline-math parens: ``\\(...\\)`` → ``$...$``
    5. Add a missing closing ``$`` when the count on a line is odd.

    Args:
        text: The text potentially containing LaTeX

    Returns:
        Text with normalised ``$...$`` LaTeX delimiters.
    """
    # 1. Backtick-enclosed LaTeX: `$...$` or `$...` -> $...$
    text = re.sub(r'`\$([^$`]+)\$`', r'$\1$', text)  # `$...$` -> $...$
    text = re.sub(r'`\$([^$`]+)`', r'$\1$', text)    # `$...` -> $...$

    # 2. Display-math double-dollars $$...$$ -> $...$
    # Use re.DOTALL so multi-line display math is handled correctly.
    text = re.sub(r'\$\$(.+?)\$\$', lambda m: '$' + m.group(1).strip() + '$',
                  text, flags=re.DOTALL)

    # 3. LaTeX \[...\] display-math brackets -> $...$
    text = re.sub(r'\\\[(.+?)\\\]', lambda m: '$' + m.group(1).strip() + '$',
                  text, flags=re.DOTALL)

    # 4. LaTeX \(...\) inline-math parens -> $...$
    text = re.sub(r'\\\((.+?)\\\)', lambda m: '$' + m.group(1).strip() + '$',
                  text, flags=re.DOTALL)

    # 5. Fix unmatched single $ delimiters (odd count per line)
    dollar_positions = [m.start() for m in re.finditer(r'\$', text)]
    if len(dollar_positions) % 2 == 1:
        lines = text.split('\n')
        fixed_lines = []
        for line in lines:
            if line.count('$') % 2 == 1:
                line = line.rstrip() + '$'
            fixed_lines.append(line)
        text = '\n'.join(fixed_lines)

    return text


def format_inline_markdown(text: str, styles=None, temp_image_files=None) -> str:
    """
    Format inline markdown elements like bold, italic, code, and LaTeX.
    Converts to ReportLab XML markup.
    LaTeX formulas are converted to images and inserted at the correct position.

    Args:
        text: The text to format
        styles: ReportLab styles (optional, needed for complex formatting)
        temp_image_files: List to track temporary image files for cleanup

    Returns:
        Either a string with XML markup, or a list of flowables if images are present

    Note: ReportLab's Paragraph class handles XML-like markup, so we don't
    escape angle brackets for bold/italic tags that we intentionally add.
    XML-escaping is applied to non-LaTeX text segments ONLY so that LaTeX
    formulas (which may legitimately contain ``<``, ``>``, ``&``) are not
    corrupted before being sent to matplotlib for rendering.
    """
    if temp_image_files is None:
        temp_image_files = []

    # Normalise all LaTeX delimiter variants to $...$ BEFORE any escaping.
    text = fix_latex_delimiters(text)

    # Check if text contains LaTeX formulas (inspect raw text, before sanitising)
    latex_pattern = r'\$([^$]+?)\$'
    has_latex = re.search(latex_pattern, text)

    if has_latex:
        # Split text by LaTeX formulas and create a mix of text and images.
        # Only the non-LaTeX segments are XML-escaped; the formula content is
        # passed as-is to matplotlib so characters like < > & render correctly.
        parts = []
        last_end = 0

        for match in re.finditer(latex_pattern, text):
            # Sanitize and format the plain-text segment before the formula.
            before_text = text[last_end:match.start()]
            if before_text:
                formatted_text = apply_markdown_formatting(sanitize_for_xml(before_text))
                parts.append(('text', formatted_text))

            # Render LaTeX formula as image (no XML escaping – raw LaTeX source).
            latex_formula = match.group(1).strip()
            image_path = render_latex_to_image(latex_formula, inline=True)

            if image_path:
                temp_image_files.append(image_path)
                # Store both image path and formula for fallback
                parts.append(('image', image_path, latex_formula))
            else:
                # Fallback to monospace if image rendering fails.
                # Sanitize the formula text so it is safe in ReportLab XML.
                safe_formula = sanitize_for_xml(latex_formula)
                parts.append(('text', f'<font name="Courier">{safe_formula}</font>'))

            last_end = match.end()

        # Sanitize and format any trailing plain-text segment.
        if last_end < len(text):
            remaining_text = text[last_end:]
            formatted_text = apply_markdown_formatting(sanitize_for_xml(remaining_text))
            parts.append(('text', formatted_text))

        # If we have images, we need to return a list of flowables
        if any(p[0] == 'image' for p in parts):
            if styles is None:
                # Can't create flowables without styles, return text only with formulas in monospace
                result = []
                for p in parts:
                    if p[0] == 'text':
                        result.append(p[1])
                    else:  # image
                        # Use the formula text (p[2]) for fallback
                        result.append(f'<font name="Courier">{sanitize_for_xml(p[2])}</font>')
                return ''.join(result)

            flowables = []
            for p in parts:
                part_type = p[0]
                if part_type == 'text':
                    part_value = p[1]
                    if part_value.strip():
                        flowables.append(Paragraph(part_value, styles['CustomBody']))
                elif part_type == 'image':
                    part_value = p[1]  # image path
                    try:
                        # Create inline image with appropriate size
                        img = Image(part_value)

                        # Get the actual image dimensions
                        actual_width = img.imageWidth
                        actual_height = img.imageHeight

                        # Scale to a reasonable inline size (0.5 inch height as base)
                        # This gives us readable math inline with text
                        target_height = 0.5 * inch

                        # Maintain aspect ratio
                        aspect_ratio = actual_width / float(actual_height)
                        img.drawHeight = target_height
                        img.drawWidth = target_height * aspect_ratio

                        # Cap maximum width to avoid overflow
                        max_width = 4 * inch
                        if img.drawWidth > max_width:
                            img.drawWidth = max_width
                            img.drawHeight = max_width / aspect_ratio

                        flowables.append(img)
                        logger.debug(f"Added image flowable: {part_value} ({img.drawWidth}, {img.drawHeight})")
                    except Exception as e:
                        logger.error(f"Failed to create image from {part_value}: {str(e)}")

            return flowables if flowables else ''
        else:
            # No images, just return formatted text
            return ''.join(p[1] for p in parts)
    else:
        # No LaTeX formulas – sanitize the whole text then apply markdown.
        return apply_markdown_formatting(sanitize_for_xml(text))


def sanitize_for_xml(text: str) -> str:
    """
    Escape special XML characters in text for ReportLab's Paragraph XML parser.
    Also strips non-printable / non-renderable Unicode characters that cause ■
    glyphs in PDF output.
    Must be called before adding any ReportLab XML markup tags.
    """
    # Strip carriage returns and other control chars (except \n which is
    # already handled by the line splitter).
    text = text.replace('\r', '')
    # Remove Unicode chars that common PDF fonts can't render
    text = re.sub(r'[\u200b-\u200f\u2028\u2029\u2060\ufeff\ufffe\uffff\u00ad]', '', text)
    # XML-escape
    text = text.replace('&', '&amp;')
    text = text.replace('<', '&lt;')
    text = text.replace('>', '&gt;')
    return text


def apply_markdown_formatting(text: str) -> str:
    """
    Apply markdown formatting (bold, italic, code) to text.
    This is separated from format_inline_markdown to handle LaTeX formulas properly.

    Processing order matters to avoid malformed XML:
    1. Inline code is extracted FIRST and replaced with placeholders so that
       characters like * and _ inside code spans are not treated as formatting.
    2. Bold/italic formatting is applied to the remaining text.
    3. Placeholders are restored with the properly wrapped <font> tags.
    """
    # --- Step 1: Extract inline code spans into placeholders ---
    code_spans = []

    def _replace_code(match):
        idx = len(code_spans)
        code_spans.append(match.group(1))
        return f"\x00CODE{idx}\x00"

    text = re.sub(r'`([^`$]+)`', _replace_code, text)

    # --- Step 2: Apply bold / italic formatting (safe now — no code content) ---
    # Bold + italic: ***text*** or ___text___
    text = re.sub(r'\*\*\*(.+?)\*\*\*', r'<b><i>\1</i></b>', text)
    text = re.sub(r'___(.+?)___', r'<b><i>\1</i></b>', text)

    # Bold: **text**
    text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)

    # Italic: *text* (not mid-word)
    text = re.sub(r'(?<!\w)\*([^*]+?)\*(?!\w)', r'<i>\1</i>', text)

    # Italic: _text_ (not inside identifiers like __init__)
    text = re.sub(r'(?<![_a-zA-Z0-9])_([^_]+?)_(?![_a-zA-Z0-9])', r'<i>\1</i>', text)

    # --- Step 3: Restore code spans with proper <font> wrapping ---
    for idx, code_content in enumerate(code_spans):
        text = text.replace(
            f"\x00CODE{idx}\x00",
            f'<font name="Courier" color="#333333">{code_content}</font>',
        )

    return text


# ==========================
# Export Service
# ==========================
class ExportService:

    @staticmethod
    def export_to_markdown(content: str, title: str) -> io.BytesIO:
        """Export content as markdown file."""
        output = io.BytesIO()
        md = f"# {title}\n\n{content}"
        output.write(md.encode("utf-8"))
        output.seek(0)
        return output

    @staticmethod
    async def export_to_pdf(content: str, title: str, watermark: bool = True) -> io.BytesIO:
        """
        Export content as PDF with markdown and LaTeX support using ReportLab.
        LaTeX formulas are rendered as images using matplotlib.
        
        Note: ReportLab has limited emoji support. Complex emojis may not render correctly.
        """
        temp_image_files = []
        try:
            output = io.BytesIO()
            
            # Create custom canvas with watermark support
            doc = SimpleDocTemplate(
                output,
                pagesize=A4,
                rightMargin=20,
                leftMargin=20,
                topMargin=20,
                bottomMargin=40 if watermark else 20,
                title=title
            )
            
            # Get custom styles
            styles = create_styles()
            
            # Build story (content elements)
            story = []
            
            # Add title
            story.append(Paragraph(sanitize_for_xml(title), styles['CustomTitle']))
            story.append(Spacer(1, 0.3*inch))
            
            # Parse markdown content and add to story
            content_elements, temp_image_files = parse_markdown_to_reportlab(content, styles)
            story.extend(content_elements)
            
            # Build PDF with custom canvas for watermark
            if watermark:
                doc.build(
                    story,
                    canvasmaker=lambda *args, **kwargs: WatermarkCanvas(
                        *args, add_watermark=True, **kwargs
                    )
                )
            else:
                doc.build(story)
            
            output.seek(0)
            logger.info(f"PDF generated successfully for '{title}' using ReportLab with LaTeX support")
            return output
            
        except Exception as e:
            logger.error(f"PDF generation failed: {str(e)}", exc_info=True)
            raise
        finally:
            # Clean up temporary image files
            for temp_file in temp_image_files:
                try:
                    if os.path.exists(temp_file):
                        os.remove(temp_file)
                        logger.debug(f"Cleaned up temporary LaTeX image: {temp_file}")
                except Exception as e:
                    logger.warning(f"Failed to clean up temporary file {temp_file}: {str(e)}")

    @staticmethod
    def export_to_docx(content: str, title: str) -> io.BytesIO:
        """Export content as DOCX file."""
        output = io.BytesIO()
        doc = Document()

        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Simple paragraph splitting - DOCX doesn't need fancy markdown parsing
        for line in content.split("\n"):
            if not line.strip():
                continue
            doc.add_paragraph(line)

        doc.save(output)
        output.seek(0)
        return output

    @staticmethod
    def export_to_txt(content: str, title: str) -> io.BytesIO:
        """Export content as plain text file with markdown stripped."""
        output = io.BytesIO()
        # Strip common markdown syntax for plain text
        text = f"{title}\n{'=' * len(title)}\n\n{content}"
        # Remove markdown bold/italic markers
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'\*(.+?)\*', r'\1', text)
        # Remove markdown heading markers
        text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
        # Remove markdown code block markers
        text = re.sub(r'```\w*\n?', '', text)
        output.write(text.encode("utf-8"))
        output.seek(0)
        return output


# Global instance
export_service = ExportService()